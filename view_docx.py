#!/usr/bin/env python3
import docx
from pathlib import Path

def main():
    # Path to the DOCX file
    file_path = Path("data") / "Customized Resume.docx"
    
    try:
        # Open the document
        doc = docx.Document(file_path)
        
        print(f"Successfully opened: {file_path}\n")
        
        # Extract and print all paragraphs
        print("Document content:")
        print("-" * 50)
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            if paragraph.text.strip():  # Only print non-empty paragraphs
                print(f"Paragraph {i}: {paragraph.text}")
        
        print("-" * 50)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        
    except Exception as e:
        print(f"Error opening the document: {e}")

if __name__ == "__main__":
    main()

